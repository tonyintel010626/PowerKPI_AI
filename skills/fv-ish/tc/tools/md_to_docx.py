#!/usr/bin/env python3
"""
Markdown TCD/TC to Word (.docx) Converter
==========================================

Converts TCD/TC markdown files to properly formatted Word documents
with Intel blue color scheme.

Usage:
  # Convert a single file
  python md_to_docx.py input.md output.docx

  # Convert with defaults (reads from tcd_output/, writes to Downloads)
  python md_to_docx.py --all
  python md_to_docx.py --tc-only

Features:
  - Intel blue color scheme (headers: #1F4E79)
  - Table formatting with blue header rows
  - Calibri font, narrow margins
  - Supports: markdown tables, headings (H1-H4), ordered/unordered lists,
    blockquotes, code blocks, inline bold/italic/code
"""
import re
import sys
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_markdown_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator lines (|---|---|)
        if re.match(r'^\|[\s\-:]+\|', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def add_formatted_run(paragraph, text, bold=False, italic=False, font_size=None, color=None):
    """Add a formatted run to a paragraph, handling markdown bold/italic inline."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    run.font.name = 'Calibri'
    return run


def add_rich_text(paragraph, text, base_size=10):
    """Parse inline markdown formatting (bold, italic, code) and add runs."""
    # Process inline formatting
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            add_formatted_run(paragraph, part[2:-2], bold=True, font_size=base_size)
        elif part.startswith('`') and part.endswith('`'):
            run = add_formatted_run(paragraph, part[1:-1], font_size=base_size)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            add_formatted_run(paragraph, part[1:-1], italic=True, font_size=base_size)
        else:
            add_formatted_run(paragraph, part, font_size=base_size)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is None:
        shading = tcPr.makeelement(qn('w:shd'), {})
        tcPr.append(shading)
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_table_to_doc(doc, table_rows):
    """Add a formatted table to the document."""
    if not table_rows or len(table_rows) < 1:
        return

    num_cols = len(table_rows[0])
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(table_rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.style = doc.styles['Normal']
                add_rich_text(p, cell_text.strip(), base_size=9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

                # Header row styling
                if i == 0:
                    set_cell_shading(cell, '1F4E79')
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def md_to_docx(md_path, docx_path):
    """Convert a markdown TCD/TC file to a Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    i = 0
    table_buffer = []
    in_table = False
    in_blockquote = False
    blockquote_buffer = []
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines (but flush buffers)
        if not stripped:
            if in_table and table_buffer:
                rows = parse_markdown_table(table_buffer)
                add_table_to_doc(doc, rows)
                table_buffer = []
                in_table = False
            if in_blockquote and blockquote_buffer:
                bq_text = ' '.join(blockquote_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(4)
                add_rich_text(p, bq_text, base_size=9)
                # Add left border color via shading
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.makeelement(qn('w:pBdr'), {})
                left_bdr = pBdr.makeelement(qn('w:left'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '12',
                    qn('w:space'): '4',
                    qn('w:color'): 'D4A017'
                })
                pBdr.append(left_bdr)
                pPr.append(pBdr)
                blockquote_buffer = []
                in_blockquote = False
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line.rstrip())
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            # Add a thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): '1F4E79'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote lines
        if stripped.startswith('>'):
            bq_content = stripped.lstrip('>').strip()
            if not in_blockquote:
                in_blockquote = True
                blockquote_buffer = []
            blockquote_buffer.append(bq_content)
            i += 1
            continue
        elif in_blockquote and blockquote_buffer:
            bq_text = ' '.join(blockquote_buffer)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(4)
            add_rich_text(p, bq_text, base_size=9)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left_bdr = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '4',
                qn('w:color'): 'D4A017'
            })
            pBdr.append(left_bdr)
            pPr.append(pBdr)
            blockquote_buffer = []
            in_blockquote = False

        # Table lines
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(stripped)
            i += 1
            continue
        elif in_table and table_buffer:
            rows = parse_markdown_table(table_buffer)
            add_table_to_doc(doc, rows)
            table_buffer = []
            in_table = False

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown bold from heading text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 2:
                p = doc.add_heading(text, level=1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 3:
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            elif level == 4:
                p = doc.add_heading(text, level=3)
            i += 1
            continue

        # Ordered list items
        ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(1.5)
            p.text = ''
            add_rich_text(p, text, base_size=10)
            i += 1
            continue

        # Unordered list items
        ul_match = re.match(r'^[-*]\s+(.*)', stripped)
        if ul_match:
            text = ul_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.5)
            p.text = ''
            add_rich_text(p, text, base_size=10)
            i += 1
            continue

        # Indented sub-items (- or number under numbered list)
        sub_match = re.match(r'^   [-*]\s+(.*)', line)
        if sub_match:
            text = sub_match.group(1)
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.left_indent = Cm(2.5)
            p.text = ''
            add_rich_text(p, text, base_size=10)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_rich_text(p, stripped, base_size=10)
        i += 1

    # Flush remaining buffers
    if in_table and table_buffer:
        rows = parse_markdown_table(table_buffer)
        add_table_to_doc(doc, rows)
    if in_blockquote and blockquote_buffer:
        bq_text = ' '.join(blockquote_buffer)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        add_rich_text(p, bq_text, base_size=9)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) == 3:
        # Generic mode: python md_to_docx.py input.md output.docx
        md_path = sys.argv[1]
        docx_path = sys.argv[2]
        if not os.path.exists(md_path):
            print(f"Error: Input file not found: {md_path}")
            sys.exit(1)
        os.makedirs(os.path.dirname(docx_path) if os.path.dirname(docx_path) else ".", exist_ok=True)
        print(f"Converting: {md_path} -> {docx_path}")
        md_to_docx(md_path, docx_path)
    elif len(sys.argv) > 1 and sys.argv[1] in ('--all', '--tc-only', '--tcd-only'):
        # Legacy batch mode for backward compatibility
        tcd_dir = r'C:\git\applications.ai.ocode.market.skills\tcd_output'
        download_dir = r'C:\Users\pgsvlab\Downloads'

        if sys.argv[1] == '--tc-only':
            files = [
                ('TC_NVU_D0i1_20260405.md', 'TC_NVU_D0i1_20260405.docx'),
                ('TC_NVU_D0i2_20260405.md', 'TC_NVU_D0i2_20260405.docx'),
            ]
        elif sys.argv[1] == '--tcd-only':
            files = [
                ('TCD_NVU_PM_20260403.md', 'TCD_NVU_PM_20260403.docx'),
                ('TCD_NVU_WakeEvent_20260403.md', 'TCD_NVU_WakeEvent_20260403.docx'),
            ]
        elif sys.argv[1] == '--all':
            files = [
                ('TCD_NVU_PM_20260403.md', 'TCD_NVU_PM_20260403.docx'),
                ('TCD_NVU_WakeEvent_20260403.md', 'TCD_NVU_WakeEvent_20260403.docx'),
                ('TC_NVU_D0i1_20260405.md', 'TC_NVU_D0i1_20260405.docx'),
                ('TC_NVU_D0i2_20260405.md', 'TC_NVU_D0i2_20260405.docx'),
            ]
        else:
            files = []

        for md_name, docx_name in files:
            md_path = os.path.join(tcd_dir, md_name)
            docx_path = os.path.join(download_dir, docx_name)
            print(f"Converting: {md_name} -> {docx_name}")
            md_to_docx(md_path, docx_path)

        print(f"\nDone! {len(files)} file(s) saved to Downloads folder.")
    else:
        print("Usage:")
        print("  python md_to_docx.py <input.md> <output.docx>   # Convert single file")
        print("  python md_to_docx.py --all                       # Convert all TCD+TC files")
        print("  python md_to_docx.py --tcd-only                  # Convert TCD files only")
        print("  python md_to_docx.py --tc-only                   # Convert TC files only")
        sys.exit(1)
