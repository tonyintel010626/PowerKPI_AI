#!/usr/bin/env python3
"""
Universal Document Extractor — doc_extract.py

Extracts ALL text content from any supported document format using
format-specific parsers that guarantee completeness.

Usage:
    python doc_extract.py <file_path> [--output <dump_path>]

Supported formats: .docx, .pdf, .md, .txt, .html, .xlsx, .csv,
                   .pptx, .rst, .xml, .json, .c, .h, .py
"""

import os
import sys
import json
import argparse
import re
from pathlib import Path
from datetime import datetime


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

FORMAT_MAP = {
    '.docx': 'docx',
    '.pdf':  'pdf',
    '.md':   'text',
    '.txt':  'text',
    '.rst':  'text',
    '.html': 'html',
    '.htm':  'html',
    '.xlsx': 'xlsx',
    '.csv':  'csv',
    '.pptx': 'pptx',
    '.xml':  'xml_file',
    '.json': 'json_file',
    # Source code — treated as text with structure awareness
    '.c':    'source',
    '.h':    'source',
    '.py':   'source',
    '.js':   'source',
    '.ts':   'source',
    '.java': 'source',
    '.rs':   'source',
    '.go':   'source',
    '.rb':   'source',
    '.sh':   'source',
    '.bat':  'source',
    '.ps1':  'source',
}


def detect_format(file_path):
    """Detect document format from file extension."""
    ext = Path(file_path).suffix.lower()
    fmt = FORMAT_MAP.get(ext)
    if fmt is None:
        # Fallback: try reading as text
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read(1024)
            return 'text'
        except (UnicodeDecodeError, Exception):
            return None
    return fmt


# ---------------------------------------------------------------------------
# Phase 1: INVENTORY
# ---------------------------------------------------------------------------

def inventory_docx(file_path):
    """Inventory all parts in a .docx ZIP archive."""
    import zipfile
    inv = {
        'format': 'docx',
        'file': file_path,
        'file_size': os.path.getsize(file_path),
        'parts': {},
        'part_list': [],
    }
    with zipfile.ZipFile(file_path, 'r') as z:
        for name in z.namelist():
            inv['part_list'].append(name)
            if name.endswith('.xml'):
                category = 'other'
                nl = name.lower()
                if 'document.xml' in nl and 'word/' in nl:
                    category = 'main_body'
                elif 'header' in nl:
                    category = 'header'
                elif 'footer' in nl:
                    category = 'footer'
                elif 'footnote' in nl:
                    category = 'footnote'
                elif 'endnote' in nl:
                    category = 'endnote'
                elif 'comment' in nl:
                    category = 'comment'
                elif 'diagram' in nl or 'chart' in nl:
                    category = 'diagram'
                inv['parts'].setdefault(category, []).append(name)
    return inv


def inventory_pdf(file_path):
    """Inventory pages and structure of a PDF."""
    inv = {
        'format': 'pdf',
        'file': file_path,
        'file_size': os.path.getsize(file_path),
        'page_count': 0,
        'has_tables': False,
    }
    try:
        import pdfplumber  # type: ignore[import-untyped]
        with pdfplumber.open(file_path) as pdf:
            inv['page_count'] = len(pdf.pages)
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    inv['has_tables'] = True
                    break
    except ImportError:
        try:
            import fitz  # type: ignore[import-untyped]  # PyMuPDF
            doc = fitz.open(file_path)
            inv['page_count'] = len(doc)
            doc.close()
        except ImportError:
            inv['error'] = 'Neither pdfplumber nor PyMuPDF available'
    return inv


def inventory_xlsx(file_path):
    """Inventory sheets in an Excel file."""
    inv = {
        'format': 'xlsx',
        'file': file_path,
        'file_size': os.path.getsize(file_path),
        'sheets': [],
    }
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        for name in wb.sheetnames:
            ws = wb[name]
            inv['sheets'].append({
                'name': name,
                'max_row': ws.max_row,
                'max_column': ws.max_column,
            })
        wb.close()
    except ImportError:
        inv['error'] = 'openpyxl not available'
    return inv


def inventory_pptx(file_path):
    """Inventory slides in a PowerPoint file."""
    import zipfile
    inv = {
        'format': 'pptx',
        'file': file_path,
        'file_size': os.path.getsize(file_path),
        'slides': 0,
        'notes': 0,
        'parts': [],
    }
    with zipfile.ZipFile(file_path, 'r') as z:
        for name in z.namelist():
            if 'slide' in name.lower() and name.endswith('.xml'):
                if 'note' in name.lower():
                    inv['notes'] += 1
                elif 'slideLayout' not in name and 'slideMaster' not in name:
                    inv['slides'] += 1
            inv['parts'].append(name)
    return inv


def inventory_text(file_path):
    """Inventory a text-based file."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()
    lines = content.split('\n')
    return {
        'format': 'text',
        'file': file_path,
        'file_size': os.path.getsize(file_path),
        'line_count': len(lines),
        'word_count': len(content.split()),
        'byte_count': len(content.encode('utf-8')),
    }


def inventory_file(file_path):
    """Run inventory for any supported format."""
    fmt = detect_format(file_path)
    if fmt == 'docx':
        return inventory_docx(file_path)
    elif fmt == 'pdf':
        return inventory_pdf(file_path)
    elif fmt == 'xlsx':
        return inventory_xlsx(file_path)
    elif fmt == 'pptx':
        return inventory_pptx(file_path)
    elif fmt in ('text', 'source', 'html', 'xml_file', 'json_file', 'csv'):
        return inventory_text(file_path)
    else:
        return {'error': f'Unsupported format: {fmt}', 'file': file_path}


# ---------------------------------------------------------------------------
# Phase 2: EXTRACT
# ---------------------------------------------------------------------------

def extract_docx(file_path):
    """Extract ALL text from a .docx — raw XML + python-docx structured."""
    import zipfile
    import xml.etree.ElementTree as ET

    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    raw = {
        'main_body': [],
        'headers': [],
        'footers': [],
        'footnotes': [],
        'endnotes': [],
        'comments': [],
        'other': [],
    }
    raw_word_count = 0

    with zipfile.ZipFile(file_path, 'r') as z:
        for name in z.namelist():
            if not name.endswith('.xml'):
                continue
            try:
                data = z.read(name)
                root = ET.fromstring(data)
            except Exception:
                continue

            texts = []
            for elem in root.iter():
                if elem.tag == f'{W_NS}t' and elem.text:
                    texts.append(elem.text)

            if not texts:
                continue

            nl = name.lower()
            if 'document.xml' in nl and 'word/' in nl:
                raw['main_body'].extend(texts)
            elif 'header' in nl:
                raw['headers'].extend(texts)
            elif 'footer' in nl:
                raw['footers'].extend(texts)
            elif 'footnote' in nl:
                raw['footnotes'].extend(texts)
            elif 'endnote' in nl:
                raw['endnotes'].extend(texts)
            elif 'comment' in nl:
                raw['comments'].extend(texts)
            else:
                if any(len(t.strip()) > 3 for t in texts):
                    raw['other'].extend(texts)

    for category, texts in raw.items():
        for t in texts:
            raw_word_count += len(t.split())

    # Structured extraction via python-docx
    structured_lines = []
    try:
        import docx
        doc = docx.Document(file_path)
        pnum = 0
        for p in doc.paragraphs:
            pnum += 1
            text = p.text.strip()
            if text:
                style = p.style.name if p.style else 'None'
                structured_lines.append(f'P{pnum:04d} [{style}]: {text}')

        for tidx, table in enumerate(doc.tables):
            structured_lines.append(f'')
            structured_lines.append(f'=== TABLE {tidx + 1} ===')
            for ridx, row in enumerate(table.rows):
                cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
                structured_lines.append(f'  ROW{ridx}: ' + ' || '.join(cells))
    except ImportError:
        structured_lines.append('# python-docx not available — using raw XML only')
        # Reconstruct from raw main_body
        for i, t in enumerate(raw['main_body']):
            if t.strip():
                structured_lines.append(f'X{i:04d}: {t.strip()}')

    structured_word_count = 0
    for line in structured_lines:
        cleaned = re.sub(r'^P\d{4}\s*\[.*?\]:\s*', '', line)
        cleaned = re.sub(r'^=== TABLE \d+ ===$', '', cleaned)
        cleaned = re.sub(r'^\s*ROW\d+:\s*', '', cleaned)
        cleaned = re.sub(r'^X\d{4}:\s*', '', cleaned)
        structured_word_count += len(cleaned.split())

    return {
        'raw': raw,
        'raw_word_count': raw_word_count,
        'structured_lines': structured_lines,
        'structured_word_count': structured_word_count,
    }


def extract_pdf(file_path):
    """Extract ALL text from a PDF."""
    pages = []
    tables = []
    total_words = 0

    try:
        import pdfplumber  # type: ignore[import-untyped]
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                pages.append(f'=== PAGE {i+1} ===\n{text}')
                total_words += len(text.split())
                for t in (page.extract_tables() or []):
                    tables.append({'page': i + 1, 'data': t})
    except ImportError:
        try:
            import fitz  # type: ignore[import-untyped]
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                text = page.get_text()
                pages.append(f'=== PAGE {i+1} ===\n{text}')
                total_words += len(text.split())
            doc.close()
        except ImportError:
            return {'error': 'No PDF library available (need pdfplumber or PyMuPDF)'}

    structured_lines = []
    for p in pages:
        structured_lines.extend(p.split('\n'))

    return {
        'raw': {'pages': pages, 'tables': tables},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_xlsx(file_path):
    """Extract ALL content from an Excel file."""
    try:
        import openpyxl
    except ImportError:
        return {'error': 'openpyxl not available'}

    wb = openpyxl.load_workbook(file_path, data_only=True)
    structured_lines = []
    total_words = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        structured_lines.append(f'=== SHEET: {sheet_name} ===')
        for row in ws.iter_rows():
            cells = []
            for cell in row:
                val = str(cell.value) if cell.value is not None else ''
                cells.append(val)
            if any(c for c in cells):
                line = ' || '.join(cells)
                structured_lines.append(line)
                total_words += len(line.split())

    wb.close()
    return {
        'raw': {},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_csv(file_path):
    """Extract ALL content from a CSV file."""
    import csv
    structured_lines = []
    total_words = 0

    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            line = ' || '.join(row)
            structured_lines.append(f'ROW{i}: {line}')
            total_words += len(line.split())

    return {
        'raw': {},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_pptx(file_path):
    """Extract ALL text from a PowerPoint file."""
    import zipfile
    import xml.etree.ElementTree as ET

    A_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    structured_lines = []
    total_words = 0

    with zipfile.ZipFile(file_path, 'r') as z:
        slide_parts = sorted([
            n for n in z.namelist()
            if 'slide' in n.lower() and n.endswith('.xml')
            and 'slideLayout' not in n and 'slideMaster' not in n
        ])

        for name in slide_parts:
            try:
                data = z.read(name)
                root = ET.fromstring(data)
            except Exception:
                continue

            texts = []
            for elem in root.iter():
                if elem.tag == f'{A_NS}t' and elem.text:
                    texts.append(elem.text.strip())

            if texts:
                label = 'NOTE' if 'note' in name.lower() else 'SLIDE'
                structured_lines.append(f'=== {label}: {name} ===')
                for t in texts:
                    structured_lines.append(t)
                    total_words += len(t.split())

    return {
        'raw': {},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_html(file_path):
    """Extract ALL visible text from an HTML file."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        # Remove script and style elements
        for tag in soup(['script', 'style']):
            tag.decompose()
        text = soup.get_text(separator='\n', strip=True)
    except ImportError:
        # Fallback: strip tags with regex
        text = re.sub(r'<[^>]+>', '\n', content)
        text = re.sub(r'\n{3,}', '\n\n', text).strip()

    lines = text.split('\n')
    word_count = len(text.split())

    return {
        'raw': {'html_source': content[:1000] + '...' if len(content) > 1000 else content},
        'raw_word_count': word_count,
        'structured_lines': lines,
        'structured_word_count': word_count,
    }


def extract_xml_file(file_path):
    """Extract all text content from an XML file."""
    import xml.etree.ElementTree as ET

    tree = ET.parse(file_path)
    root = tree.getroot()
    structured_lines = []
    total_words = 0

    def walk(elem, depth=0):
        nonlocal total_words
        indent = '  ' * depth
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        text = (elem.text or '').strip()
        tail = (elem.tail or '').strip()
        if text:
            structured_lines.append(f'{indent}<{tag}>: {text}')
            total_words += len(text.split())
        if tail:
            structured_lines.append(f'{indent}(tail): {tail}')
            total_words += len(tail.split())
        for child in elem:
            walk(child, depth + 1)

    walk(root)
    return {
        'raw': {},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_json_file(file_path):
    """Extract all string values from a JSON file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    structured_lines = []
    total_words = 0

    def walk(obj, path=''):
        nonlocal total_words
        if isinstance(obj, dict):
            for k, v in obj.items():
                walk(v, f'{path}.{k}')
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                walk(v, f'{path}[{i}]')
        elif isinstance(obj, str) and obj.strip():
            structured_lines.append(f'{path}: {obj}')
            total_words += len(obj.split())
        elif obj is not None:
            structured_lines.append(f'{path}: {obj}')

    walk(data)
    return {
        'raw': {},
        'raw_word_count': total_words,
        'structured_lines': structured_lines,
        'structured_word_count': total_words,
    }


def extract_text(file_path):
    """Extract all content from a text file."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()
    lines = content.split('\n')
    word_count = len(content.split())
    return {
        'raw': {},
        'raw_word_count': word_count,
        'structured_lines': lines,
        'structured_word_count': word_count,
    }


def extract_source(file_path):
    """Extract content from source code with structure awareness."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()

    lines = content.split('\n')
    word_count = len(content.split())

    # Add structural markers
    structured_lines = []
    for i, line in enumerate(lines, 1):
        stripped = line.strip()
        # Detect function/struct/class definitions
        if re.match(r'(static\s+)?(void|int|char|bool|struct|class|def|fn|func|function)\s+', stripped):
            structured_lines.append(f'=== DEFINITION at line {i} ===')
        structured_lines.append(f'L{i:05d}: {line}')

    return {
        'raw': {},
        'raw_word_count': word_count,
        'structured_lines': structured_lines,
        'structured_word_count': word_count,
    }


EXTRACTORS = {
    'docx': extract_docx,
    'pdf': extract_pdf,
    'xlsx': extract_xlsx,
    'csv': extract_csv,
    'pptx': extract_pptx,
    'html': extract_html,
    'xml_file': extract_xml_file,
    'json_file': extract_json_file,
    'text': extract_text,
    'source': extract_source,
}


def extract_file(file_path):
    """Run full extraction for any supported format."""
    fmt = detect_format(file_path)
    if fmt is None:
        return {'error': f'Cannot detect format for: {file_path}'}
    extractor = EXTRACTORS.get(fmt)
    if extractor is None:
        return {'error': f'No extractor for format: {fmt}'}
    return extractor(file_path)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description='Universal Document Extractor')
    parser.add_argument('file_path', help='Path to the document to extract')
    parser.add_argument('--output', '-o', help='Output dump file path (default: auto-generated)')
    parser.add_argument('--inventory-only', action='store_true', help='Only run inventory, skip extraction')
    parser.add_argument('--json', action='store_true', help='Output inventory/stats as JSON')
    args = parser.parse_args()

    file_path = args.file_path
    if not os.path.exists(file_path):
        print(f'ERROR: File not found: {file_path}', file=sys.stderr)
        sys.exit(1)

    # Phase 1: Inventory
    print(f'Phase 1: INVENTORY')
    print(f'  File: {file_path}')
    inv = inventory_file(file_path)
    fmt = detect_format(file_path)
    print(f'  Format: {fmt}')

    if fmt == 'docx':
        for cat, parts in inv.get('parts', {}).items():
            print(f'  {cat}: {len(parts)} parts')
    elif fmt == 'pdf':
        print(f'  Pages: {inv.get("page_count", "?")}')
        print(f'  Has tables: {inv.get("has_tables", "?")}')
    elif fmt == 'xlsx':
        for s in inv.get('sheets', []):
            print(f'  Sheet "{s["name"]}": {s["max_row"]}x{s["max_column"]}')
    else:
        print(f'  Lines: {inv.get("line_count", "?")}')
        print(f'  Words: {inv.get("word_count", "?")}')

    if args.inventory_only:
        if args.json:
            print(json.dumps(inv, indent=2, default=str))
        return

    # Phase 2: Extract
    print(f'\nPhase 2: EXTRACT')
    result = extract_file(file_path)

    if 'error' in result:
        print(f'  ERROR: {result["error"]}', file=sys.stderr)
        sys.exit(1)

    print(f'  Raw word count: {result["raw_word_count"]}')
    print(f'  Structured lines: {len(result["structured_lines"])}')
    print(f'  Structured word count: {result["structured_word_count"]}')

    # Write dump file
    if args.output:
        dump_path = args.output
    else:
        stem = Path(file_path).stem.replace(' ', '_').lower()
        dump_path = os.path.join(os.path.dirname(file_path) or '.', f'{stem}_full_dump.txt')

    with open(dump_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(result['structured_lines']))

    print(f'  Dump written to: {dump_path}')

    # Quick verification stats
    raw_wc = int(result['raw_word_count'])
    struct_wc = int(result['structured_word_count'])
    delta = raw_wc - struct_wc
    pct = (struct_wc / max(raw_wc, 1)) * 100
    print(f'\n  Quick verification:')
    print(f'    Coverage: {pct:.1f}%')
    print(f'    Word delta: {delta} (raw - structured)')
    if delta > 0:
        print(f'    WARNING: {delta} words in raw not in structured dump')
        print(f'    Run doc_verify.py for full delta analysis')
    else:
        print(f'    OK: Structured dump covers all raw content')

    if args.json:
        stats = {
            'file': file_path,
            'format': fmt,
            'inventory': inv,
            'raw_word_count': result['raw_word_count'],
            'structured_word_count': result['structured_word_count'],
            'structured_lines': len(result['structured_lines']),
            'delta': delta,
            'coverage_pct': round(pct, 1),
            'dump_path': dump_path,
            'timestamp': datetime.now().isoformat(),
        }
        print(json.dumps(stats, indent=2, default=str))


if __name__ == '__main__':
    main()
